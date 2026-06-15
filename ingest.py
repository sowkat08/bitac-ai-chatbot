import os
import hashlib
import time
import pandas as pd

from pinecone import Pinecone
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_cohere import CohereEmbeddings
from cohere.errors.too_many_requests_error import TooManyRequestsError

# ================= CONFIG =================
INDEX_NAME = os.getenv("INDEX_NAME", "bitac-chatbot")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
COHERE_API_KEY = os.getenv("COHERE_API_KEY")

if not PINECONE_API_KEY or not COHERE_API_KEY:
    raise ValueError("Missing API keys in Environment Variables!")

pc = Pinecone(api_key=PINECONE_API_KEY)
index = pc.Index(INDEX_NAME)

def uid(text, src):
    return hashlib.md5((text + src).encode()).hexdigest()

# ================= SMART CHECKING FUNCTION =================
def is_file_already_ingested(file_path_or_url):
    try:
        dummy_vector = [0.1] * 1024
        results = index.query(
            vector=dummy_vector,
            filter={"source": {"$eq": file_path_or_url}},
            top_k=1,
            include_metadata=False
        )
        return len(results.get('matches', [])) > 0
    except Exception as e:
        print(f"⚠️ Pinecone Query Warning for {file_path_or_url}: {e}")
        return False

# ================= ১. গিটহাব ফোল্ডারের ফাইল ট্র্যাক =================
current_active_files = set()
if os.path.exists("bitac_files"):
    for root, _, files in os.walk("bitac_files"):
        for f in files:
            full_path = os.path.join(root, f)
            normalized_path = full_path.replace("\\\\", "/").replace("\\", "/")
            current_active_files.add(normalized_path)

print(f"📁 গিটহাব ফোল্ডারে বর্তমানে মোট একটিভ ফাইল আছে: {len(current_active_files)} টি")

# ================= ২. 🔥 ফিক্সড অটোমেটিক হার্ড ডিলিট (পাইনকোন সিঙ্ক) =================
print("\n🔄 ডাটাবেজ অটো-সিঙ্ক ও ক্লিনআপ রান হচ্ছে...")

try:
    ingested_sources_in_pinecone = set()
    
    # পাইনকোন থেকে মেটাডাটা ট্র্যাকিং ফেচ পদ্ধতি
    dummy_vector = [0.1] * 1024
    query_res = index.query(
        vector=dummy_vector,
        top_k=100, 
        include_metadata=True
    )
    
    for match in query_res.get('matches', []):
        src = match.get('metadata', {}).get('source')
        if src:
            ingested_sources_in_pinecone.add(src.replace("\\\\", "/").replace("\\", "/"))

    # ওল্ড ফাইল আইডেন্টিফিকেশন লজিক
    files_to_delete = ingested_sources_in_pinecone - current_active_files
    
    if files_to_delete:
        print(f"🧹 গিটহাবে নেই কিন্তু ডাটাবেজে ওল্ড ডাটা আছে: {len(files_to_delete)} টি")
        for old_file in files_to_delete:
            print(f"🗑️ Deleting from Pinecone: {old_file}")
            index.delete(filter={"source": {"$eq": old_file}})
        print("✅ ওল্ড ফাইলগুলো পাইনকোন থেকে সফলভাবে মুছে ফেলা হয়েছে!")
    else:
        print("✨ প্রথম ধাপে কোনো ওল্ড ফাইল ডিলিট করার প্রয়োজন হয়নি।")

except Exception as e:
    print(f"⚠️ সিঙ্ক সতর্কতা: {e}")

# ================= ৩. নতুন ফাইল রিড ও অ্যাড করার লজিক =================
docs = []

if os.path.exists("bitac_files"):
    for root, _, files in os.walk("bitac_files"):
        for f in files:
            path = os.path.join(root, f)
            normalized_path_for_source = path.replace("\\\\", "/").replace("\\", "/")

            if is_file_already_ingested(normalized_path_for_source):
                print(f"Skip (Already Ingested): {normalized_path_for_source}")
                continue

            print(f"📖 Reading New File: {normalized_path_for_source}")
            try:
                if f.endswith(".txt"):
                    from langchain_community.document_loaders import TextLoader
                    loaded_docs = TextLoader(path, encoding='utf-8').load()
                    for d in loaded_docs:
                        d.metadata["source"] = normalized_path_for_source
                    docs += loaded_docs
                    
                elif f.endswith(".pdf"):
                    import pdfplumber
                    try:
                        with pdfplumber.open(path) as pdf:
                            page_count = 0
                            for page in pdf.pages:
                                text = page.extract_text() or ""
                                tables = page.extract_tables()
                                
                                if tables:
                                    text += "\n\n--- [Table Data] ---"
                                    for table in tables:
                                        for row in table:
                                            clean_row = [str(cell).strip() for cell in row if cell is not None]
                                            if clean_row:
                                                text += "\n" + " | ".join(clean_row)
                                
                                if text.strip():
                                    docs.append(Document(page_content=text, metadata={"source": normalized_path_for_source}))
                                
                                page_count += 1
                                if page_count > 100:
                                    break
                    except Exception as e:
                        print(f"❌ PDF error ({f}):", e)
                        
                elif f.endswith(".docx"):
                    import docx as py_docx
                    try:
                        doc_obj = py_docx.Document(path)
                        full_text = []
                        for para in doc_obj.paragraphs:
                            if para.text.strip():
                                full_text.append(para.text)
                        for table in doc_obj.tables:
                            full_text.append("\n--- [Table Data] ---")
                            for row in table.rows:
                                row_data = [cell.text.strip() for cell in row.cells]
                                if any(row_data): 
                                    full_text.append(" | ".join(row_data))
                        text_content = "\n".join(full_text)
                        if text_content.strip():
                            docs.append(Document(page_content=text_content, metadata={"source": normalized_path_for_source}))
                    except Exception as e:
                        print(f"❌ Word error ({f}):", e)
                        
                elif f.endswith(".xlsx"):
                    df = pd.read_excel(path)
                    docs.append(Document(page_content=df.astype(str).to_string(), metadata={"source": normalized_path_for_source}))
                    
            except Exception as e:
                print(f"❌ General File error ({f}):", e)

# ================= ৪. SPLIT, EMBED & UPLOAD =================
if not docs:
    print("✅ কোনো নতুন ফাইল যোগ করার প্রয়োজন নেই। ডাটাবেজ আপ-টু-ডেট!")
else:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(docs)

    print(f"🧠 Generating Embeddings & Uploading {len(chunks)} chunks...")
    embeddings = CohereEmbeddings(
        model="embed-multilingual-v3.0",
        cohere_api_key=COHERE_API_KEY
    )

    batch_size = 30  
    total_batches = (len(chunks) + batch_size - 1) // batch_size
    total_uploaded = 0

    for i in range(0, len(chunks), batch_size):
        batch_chunks = chunks[i:i+batch_size]
        batch_texts = [c.page_content for c in batch_chunks]
        current_batch = (i // batch_size) + 1
        
        print(f"⏳ Processing batch {current_batch} of {total_batches}...")
        
        while True:
            try:
                batch_vectors = embeddings.embed_documents(batch_texts)
                
                pinecone_upserts = []
                for j in range(len(batch_texts)):
                    src_metadata = batch_chunks[j].metadata.get("source", "file")
                    vid = uid(batch_texts[j], src_metadata)
                    
                    pinecone_upserts.append((
                        vid,
                        batch_vectors[j],
                        {
                            "text": batch_texts[j],
                            "source": src_metadata
                        }
                    ))
                
                index.upsert(vectors=pinecone_upserts)
                total_uploaded += len(pinecone_upserts)
                print(f"✅ Batch {current_batch} সফলভাবে পাইনকোনে অ্যাড হয়েছে।")
                
                time.sleep(2) 
                break 
                
            except TooManyRequestsError:
                print("\n⚠️ Cohere Rate Limit! ৬০ সেকেন্ড অপেক্ষা করছি...")
                time.sleep(60)
            except Exception as e:
                print(f"❌ Unexpected error: {e}")
                raise e

    print(f"\n🎉 সফলভাবে ইনজেস্ট সম্পন্ন! নতুন ভেক্টর সংখ্যা: {total_uploaded}")
    print("🎯 পাইনকোন ডাটাবেজ এবং গিটহাব ফোল্ডার এখন ১০০% সিঙ্কড!")
